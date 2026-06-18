# -*- coding: utf-8 -*-
"""Извлечение текста из загруженных файлов."""
from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
import openpyxl
import xlrd


ALLOWED_EXT = {".txt", ".md", ".docx", ".doc", ".xlsx", ".xls"}


def _read_text_bytes(data: bytes, filename: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _parse_txt_md(data: bytes) -> str:
    return _read_text_bytes(data, "")


def _parse_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_doc_fallback(data: bytes) -> str:
    """Грубое извлечение текста из бинарного .doc (без Word)."""
    # UTF-16LE фрагменты и ASCII-последовательности
    chunks: list[str] = []
    for m in re.finditer(rb"[\x20-\x7e\u0400-\u04ff]{4,}", data):
        try:
            s = m.group().decode("utf-8", errors="ignore")
            if len(s) >= 4 and not s.startswith("http"):
                chunks.append(s)
        except Exception:
            pass
    # UTF-16
    try:
        u16 = data.decode("utf-16-le", errors="ignore")
        words = re.findall(r"[\w\u0400-\u04ff][\w\u0400-\u04ff\s.,;:!?\-]{3,}", u16)
        chunks.extend(w for w in words if len(w.strip()) > 4)
    except Exception:
        pass
    text = "\n".join(dict.fromkeys(chunks))
    if len(text) < 80:
        raise ValueError(
            "Не удалось извлечь текст из .doc. Сохраните файл как .docx и загрузите снова."
        )
    return text


def _parse_xlsx(data: bytes) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in wb.worksheets:
        lines.append(f"## Лист: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append(" | ".join(cells))
    wb.close()
    return "\n".join(lines)


def _parse_xls(data: bytes) -> str:
    book = xlrd.open_workbook(file_contents=data)
    lines: list[str] = []
    for sheet in book.sheets():
        lines.append(f"## Лист: {sheet.name}")
        for rx in range(sheet.nrows):
            cells = [
                str(sheet.cell_value(rx, cx)).strip()
                for cx in range(sheet.ncols)
                if str(sheet.cell_value(rx, cx)).strip()
            ]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def parse_file(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXT:
        raise ValueError(f"Формат {ext} не поддерживается. Допустимо: {', '.join(sorted(ALLOWED_EXT))}")
    if ext in (".txt", ".md"):
        return _parse_txt_md(data)
    if ext == ".docx":
        return _parse_docx(data)
    if ext == ".doc":
        return _parse_doc_fallback(data)
    if ext == ".xlsx":
        return _parse_xlsx(data)
    if ext == ".xls":
        return _parse_xls(data)
    raise ValueError(f"Неизвестный формат: {ext}")


def synthesize_documents(brief: str, concept: str, details: str) -> str:
    """Склеивает три документа в единый контекст для ИИ."""
    parts = [
        "# БРИФ ПРОЕКТА\n",
        brief.strip() or "(пусто)",
        "\n\n# КОНЦЕПЦИЯ ПРОЕКТА\n",
        concept.strip() or "(пусто)",
        "\n\n# ДОПОЛНИТЕЛЬНЫЕ ДЕТАЛИ\n",
        details.strip() or "(пусто)",
    ]
    return "".join(parts)
